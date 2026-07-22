#!/usr/bin/env python3
"""供应链产品短视频内容包 — DOCX 组装脚本

读取工作流各阶段输出的 JSON，组装为格式化的 Word 文档，保存到桌面。

用法:
    python assemble_docx.py <input.json>
    python assemble_docx.py <input.json> --output <path.docx>
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json, sys, os
from datetime import datetime

DESKTOP = r"C:\Users\19821\Desktop"


# ── 工具函数 (复用 generate_docs.py 模式) ──────────────────

def add_run_with_font(paragraph, text, font_name_cn='宋体', font_name_en='Times New Roman', size=Pt(12), bold=False, color=None):
    run = paragraph.add_run(text)
    run.font.size = size
    run.font.bold = bold
    run.font.name = font_name_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    if color:
        run.font.color.rgb = color
    return run


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.5):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_title(doc, text, size=Pt(22), font_cn='黑体'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12, line_spacing=1.5)
    add_run_with_font(p, text, font_name_cn=font_cn, size=size, bold=True)


def add_subtitle(doc, text, size=Pt(10.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12, line_spacing=1.5)
    add_run_with_font(p, text, font_name_cn='楷体', size=size)


def add_section_title(doc, text, size=Pt(15)):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=18, after=8, line_spacing=1.5)
    add_run_with_font(p, text, font_name_cn='黑体', size=size, bold=True)


def add_subsection_title(doc, text, size=Pt(12)):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=4, line_spacing=1.5)
    add_run_with_font(p, text, font_name_cn='楷体', size=size, bold=True)


def add_body(doc, text, indent=True, size=Pt(12)):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=3, line_spacing=1.5)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    add_run_with_font(p, text, font_name_cn='仿宋', size=size)


def add_body_no_indent(doc, text, size=Pt(12)):
    add_body(doc, text, indent=False, size=size)


def add_bold_body(doc, text, size=Pt(12)):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=3, line_spacing=1.5)
    p.paragraph_format.first_line_indent = Pt(24)
    add_run_with_font(p, text, font_name_cn='黑体', size=size, bold=True)


def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_font(p, header, font_name_cn='黑体', size=Pt(10.5), bold=True)

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c != len(row)-1 else WD_ALIGN_PARAGRAPH.LEFT
            add_run_with_font(p, str(val), font_name_cn='宋体', size=Pt(10.5))

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_separator(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_with_font(p, '— — —', font_name_cn='宋体', size=Pt(9), color=RGBColor(0xAA, 0xAA, 0xAA))


# ── 各 Section 渲染 ──────────────────────────────────────

def render_cover(doc, data):
    """封面"""
    for _ in range(4):
        doc.add_paragraph()

    add_title(doc, data.get('product_name', '未命名产品'), size=Pt(26))
    doc.add_paragraph()
    add_subtitle(doc, '供应链产品短视频内容包', size=Pt(14))

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_with_font(p, f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}", font_name_cn='楷体', size=Pt(11))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tier_label = '精品 · AI数字人路线' if data.get('tier') == 'premium' else '普通品 · 漫剧动画路线'
    add_run_with_font(p, f'内容路线：{tier_label}', font_name_cn='楷体', size=Pt(11))

    doc.add_page_break()


def render_section1_summary(doc, data):
    """Section 1 — 执行摘要"""
    add_section_title(doc, '一、执行摘要')

    product = data.get('product', {})
    add_table(doc,
        ['项目', '内容'],
        [
            ['产品名称', product.get('product_name', '')],
            ['产品品类', product.get('category', '')],
            ['产地', product.get('origin', '')],
            ['价格区间', product.get('price_range', '')],
            ['产品档次', '精品 (AI数字人)' if data.get('tier') == 'premium' else '普通品 (漫剧动画)'],
            ['档次评分', str(data.get('tier_score', '-'))],
            ['推荐平台排序', ', '.join(product.get('target_platforms', []))],
            ['目标人群', product.get('target_audience', '')],
        ],
        col_widths=[3.5, 12.0]
    )

    add_body_no_indent(doc, '')
    add_bold_body(doc, '核心策略建议:')
    strategy = data.get('strategy_summary', '')
    add_body(doc, strategy)


def render_section2_product(doc, data):
    """Section 2 — 产品档案"""
    add_section_title(doc, '二、产品档案')

    product = data.get('product', {})

    # 基本信息表
    add_subsection_title(doc, '2.1  基本信息')
    add_table(doc,
        ['字段', '内容'],
        [
            ['产品名称', product.get('product_name', '')],
            ['品类', product.get('category', '')],
            ['产地', product.get('origin', '')],
            ['价格区间', product.get('price_range', '')],
            ['品牌名称', product.get('brand_name', '-')],
            ['产量规模', product.get('production_scale', '-')],
        ],
        col_widths=[3.5, 12.0]
    )

    # 卖点
    add_subsection_title(doc, '2.2  核心卖点')
    selling_points = product.get('selling_points', [])
    if isinstance(selling_points, str):
        selling_points = [s.strip() for s in selling_points.split('\n') if s.strip()]
    for i, sp in enumerate(selling_points, 1):
        add_body(doc, f'{i}. {sp}')

    # 产地故事
    add_subsection_title(doc, '2.3  产地/品牌故事')
    add_body(doc, product.get('story', '（未提供）'))

    # 认证
    add_subsection_title(doc, '2.4  认证资质')
    certs = product.get('certifications', [])
    if isinstance(certs, str):
        certs = [c.strip() for c in certs.split('\n') if c.strip()]
    if certs:
        add_body(doc, '、'.join(certs))
    else:
        add_body(doc, '（未填写认证信息）')


def render_section3_market(doc, data):
    """Section 3 — 市场调研"""
    add_section_title(doc, '三、市场调研')

    # 竞品分析
    add_subsection_title(doc, '3.1  竞品分析摘要')
    competitor = data.get('competitor_research', '')
    if competitor:
        add_body(doc, competitor)
    else:
        add_body(doc, '（竞品分析未完成或数据不可用）')

    add_separator(doc)

    # 趋势
    add_subsection_title(doc, '3.2  平台趋势')
    trends = data.get('trend_research', '')
    if trends:
        add_body(doc, trends)
    else:
        add_body(doc, '（趋势分析未完成或数据不可用）')

    add_separator(doc)

    # 人群
    add_subsection_title(doc, '3.3  人群洞察')
    audience = data.get('audience_research', '')
    if audience:
        add_body(doc, audience)
    else:
        add_body(doc, '（人群分析未完成或数据不可用）')


def render_section4_scripts(doc, data):
    """Section 4 — 各平台内容包"""
    add_section_title(doc, '四、各平台内容包')

    scripts = data.get('scripts', {})
    platform_order = ['douyin', 'xiaohongshu', 'wechat', 'kuaishou']
    platform_names = {
        'douyin': '抖音',
        'xiaohongshu': '小红书',
        'wechat': '视频号',
        'kuaishou': '快手',
    }

    for pk in platform_order:
        if pk not in scripts:
            continue
        content = scripts[pk]
        if not content:
            continue

        add_subsection_title(doc, f'4.{platform_order.index(pk)+1}  {platform_names[pk]}内容包')

        # 脚本内容可能很长，分行展示
        if isinstance(content, str):
            for line in content.split('\n'):
                if line.strip():
                    add_body(doc, line, indent=False)
        elif isinstance(content, dict):
            for k, v in content.items():
                add_bold_body(doc, str(k))
                add_body(doc, str(v), indent=False)

        add_separator(doc)


def render_section5_visual(doc, data):
    """Section 5 — 视觉指导"""
    add_section_title(doc, '五、视觉指导方案')

    visual = data.get('visual_direction', '')
    if visual:
        if isinstance(visual, str):
            for line in visual.split('\n'):
                if line.strip():
                    add_body(doc, line, indent=False)
        else:
            add_body(doc, str(visual))
    else:
        add_body(doc, '（视觉指导未完成或数据不可用）')


def render_section6_strategy(doc, data):
    """Section 6 — 发布策略"""
    add_section_title(doc, '六、发布策略')

    # 标签策略
    add_subsection_title(doc, '6.1  标签策略')
    hashtags = data.get('hashtag_strategy', '')
    if hashtags:
        if isinstance(hashtags, str):
            for line in hashtags.split('\n'):
                if line.strip():
                    add_body(doc, line, indent=False)
        else:
            add_body(doc, str(hashtags))
    else:
        add_body(doc, '（标签策略未完成或数据不可用）')

    add_separator(doc)

    # 发布日历
    add_subsection_title(doc, '6.2  建议发布节奏')
    add_table(doc,
        ['时间', '平台', '内容主题', '备注'],
        [
            ['第1天', '抖音 + 小红书', '产品首发', '同步发布，测试反馈'],
            ['第2天', '视频号', '产地故事版', '差异化内容'],
            ['第3天', '抖音', '第二种内容角度', '根据第1天数据调整'],
            ['第4-5天', '全平台', '互动维护', '评论回复、二次传播'],
            ['第7天', '全平台', '复盘优化', '根据首周数据调整第二周策略'],
        ],
        col_widths=[2.0, 4.0, 5.0, 5.0]
    )

    add_body(doc, '提示：以上为建议节奏模板，实际发布应根据内容准备情况和目标平台的最佳发布时间灵活调整。')


def render_section7_appendix(doc, data):
    """Section 7 — 附录"""
    add_section_title(doc, '七、附录')

    add_subsection_title(doc, '7.1  工作流元数据')
    add_table(doc,
        ['项目', '内容'],
        [
            ['生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
            ['工作流版本', '1.0.0 MVP'],
            ['产品档次', '精品' if data.get('tier') == 'premium' else '普通品'],
            ['内容路线', 'AI数字人' if data.get('tier') == 'premium' else '漫剧动画'],
        ],
        col_widths=[3.5, 12.0]
    )

    add_subsection_title(doc, '7.2  推荐工具清单')
    add_table(doc,
        ['用途', '推荐工具', '说明'],
        [
            ['视频剪辑', '剪映专业版', '免费，中文优化，模板丰富'],
            ['AI数字人', 'HeyGen / 商汤如影', '精品路线首选，效果最好'],
            ['漫画制作', '来画 / Fliki', '普通品路线，动画+配音一站式'],
            ['封面设计', 'Canva / 醒图', '小红书封面必备'],
            ['数据分析', '新抖 / 千瓜', '各平台数据分析'],
        ],
        col_widths=[3.0, 4.5, 8.5]
    )

    add_body_no_indent(doc, '')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=6, line_spacing=1.5)
    add_run_with_font(p, '— 内容包结束 —', font_name_cn='楷体', size=Pt(10.5), color=RGBColor(0x99, 0x99, 0x99))


# ── 主流程 ──────────────────────────────────────────────

def assemble(data: dict, output_path: str = None):
    """组装完整 DOCX"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # 渲染各 Section
    render_cover(doc, data)
    render_section1_summary(doc, data)
    render_section2_product(doc, data)
    render_section3_market(doc, data)
    doc.add_page_break()
    render_section4_scripts(doc, data)
    doc.add_page_break()
    render_section5_visual(doc, data)
    render_section6_strategy(doc, data)
    render_section7_appendix(doc, data)

    # 保存
    if not output_path:
        product_name = data.get('product', {}).get('product_name', '未知产品')
        date_str = datetime.now().strftime('%Y%m%d')
        output_path = os.path.join(DESKTOP, f'{product_name}_视频内容包_{date_str}.docx')

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python assemble_docx.py <input.json> [--output <path.docx>]')
        print('示例: python assemble_docx.py product_data.json')
        print('      python assemble_docx.py product_data.json --output C:/Users/19821/Desktop/output.docx')
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = None

    # 解析 --output
    for i, arg in enumerate(sys.argv):
        if arg == '--output' and i + 1 < len(sys.argv):
            output_path = sys.argv[i + 1]
        elif arg == '-o' and i + 1 < len(sys.argv):
            output_path = sys.argv[i + 1]

    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    result_path = assemble(data, output_path)
    print(f'OK: {result_path}')
